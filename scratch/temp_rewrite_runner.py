import docx
import os
import sys
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from aligned_sections_snippet import (
    build_employee_utilities,
    build_warehouse_sections,
    build_procurement_sections,
    build_approver_sections,
)

doc = docx.Document("docs/test_full_rewrite.docx")

# Helper function to set table borders XML
def set_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 1/8 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D3D3D3')  # Light grey
        tblBorders.append(border)
    tblPr.append(tblBorders)

# Helper function to set cell margins/padding
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# 1. Update FAQ Q2
faq_q2_found = False
for idx, p in enumerate(doc.paragraphs):
    if 'Q2: How can payroll be recalculated' in p.text:
        p.text = "Q2: How can a rejected PR1 be resubmitted for approval?"
        doc.paragraphs[idx + 1].text = "A: If an approver requests revision or rejects a PR1, the requisitioner can navigate to My Requests, open the requisition, make the necessary corrections, and click Submit PR1 again."
        faq_q2_found = True
        print("Updated FAQ Q2.")
        break

# 2. Update FAQ Q5
faq_q5_found = False
for idx, p in enumerate(doc.paragraphs):
    if 'Q5: How do clients receive invoice notifications?' in p.text:
        p.text = "Q5: How do suppliers receive notifications of new RFQs?"
        doc.paragraphs[idx + 1].text = "A: Suppliers receive automatic email alerts containing their unique quote link when an RFQ is opened. They can also view pending bids in the Quotations panel of their portal."
        faq_q5_found = True
        print("Updated FAQ Q5.")
        break

# 3. Update Appendix B
app_b_idx = -1
for idx, p in enumerate(doc.paragraphs):
    if 'Appendix B – Notification Triggers' in p.text:
        app_b_idx = idx
        break

if app_b_idx != -1:
    triggers = [
        "Purchase Request (PR1) submitted for validation",
        "Substitute item decision submitted by requisitioner",
        "RFQ invitation sent to accredited suppliers",
        "Supplier quote submitted for RFQ canvassing",
        "Purchase Order (PO) routed for signatory approval",
        "Goods Receipt Note (GRN) closed by warehouse"
    ]
    # Replace next 6 bullet paragraphs
    for i, trigger in enumerate(triggers):
        doc.paragraphs[app_b_idx + 2 + i].text = trigger
    print("Updated Appendix B notification triggers.")

# 4. Update Appendix F
app_f_idx = -1
for idx, p in enumerate(doc.paragraphs):
    if 'Appendix F – Role Permission Examples' in p.text:
        app_f_idx = idx
        break

if app_f_idx != -1:
    doc.paragraphs[app_f_idx + 1].text = (
        "Admin: Access to system configuration, user profiles, workflow configuration, and full audit logs.\n"
        "Employee / Requestor: Requisition creation, tracking, substitute reviews, and delivery monitoring.\n"
        "Warehouse: PR1 stock validation, Goods Receipt Note (GRN) creation/receipt, and delivery tracking.\n"
        "Procurement: RFQ/canvassing management, supplier accreditation, product reviews, and PO generation.\n"
        "Approver: Sign-off authority on PR1s, PR2s, and POs with approval history review.\n"
        "Supplier: Accreditation registration, product catalog submissions, RFQ bidding, and delivery scheduling.\n"
        "TSQA: Raw Sample Evaluation (RSE) inspections, testing logs, and product approval verdicts."
    )
    print("Updated Appendix F role permissions.")

# 5. Update Table 42 Row 5
table_42 = doc.tables[42]
t42_found = False
for row in table_42.rows:
    if 'Incorrect attendance records' in row.cells[0].text:
        row.cells[0].text = "Rejected supplier quotation in RFQ"
        row.cells[1].text = "Item requires TSQA verification first"
        row.cells[2].text = "Navigate to Product Review, verify the supplier's product status, or wait for the TSQA evaluation results."
        t42_found = True
        print("Updated Table 42 Row 5.")
        break

# 6. Update and Sort Table 43 (Glossary)
table_43 = doc.tables[43]
replacements = {
    "Attendance Module": ("Purchase Request (PR1)", "Initial internal requisition document submitted by an employee to request goods or services."),
    "Client Dashboard": ("Purchase Memo (PR2)", "Level 2 procurement request generated after canvassing suppliers, showing selected quotes and bid comparisons."),
    "Payslip": ("Request for Quotation (RFQ)", "Document sent to accredited suppliers to invite bids and pricing proposals for requested items."),
    "Project Module": ("Purchase Order (PO)", "Official commercial document issued by the buyer to a supplier, committing to pay for specified goods/services."),
    "Variable Type": ("Goods Receipt Note (GRN)", "Receipt issued by the warehouse confirming delivery details, counted quantities, and accepted/rejected items."),
    "Variable Method": ("Stock on Hand (SOH)", "The current physical inventory quantity of an item available in the warehouse.")
}

raw_data = []
for r_idx in range(1, len(table_43.rows)):
    row = table_43.rows[r_idx]
    term = row.cells[0].text.strip()
    definition = row.cells[1].text.strip()
    if term in replacements:
        term, definition = replacements[term]
    elif term == "Employee Dashboard":
        definition = "The interface for employees to view requisition status, manage substitute decisions, and track deliveries."
    elif term == "In-App Notification":
        definition = "Alerts shown within the procurement system interface, typically for time-sensitive actions or updates."
    elif term == "Module":
        definition = "A distinct feature set of the procurement system, such as Requisitions, Canvassing, or Purchase Orders."
    elif term == "SMTP (Simple Mail Transfer Protocol)":
        definition = "A protocol used by the system to send outgoing emails, such as notifications or invitations."
    raw_data.append((term, definition))

extra_terms = [
    ("Supplier Accreditation", "The compliance evaluation process that a supplier must pass before they can be invited to RFQs."),
    ("Technical Standards & Quality Assurance (TSQA)", "Department responsible for conducting quality inspections and evaluations on supplier product samples."),
    ("Raw Sample Evaluation (RSE)", "Technical inspection workflow conducted by TSQA staff to verify if raw materials meet quality requirements.")
]
raw_data.extend(extra_terms)
raw_data.sort(key=lambda x: x[0].lower())

current_rows = len(table_43.rows)
target_rows = len(raw_data) + 1

if current_rows < target_rows:
    for _ in range(target_rows - current_rows):
        table_43.add_row()
elif current_rows > target_rows:
    tbl = table_43._tbl
    for _ in range(current_rows - target_rows):
        tbl.remove(tbl.tr_lst[-1])

for idx, (term, definition) in enumerate(raw_data):
    row = table_43.rows[idx + 1]
    row.cells[0].text = term
    row.cells[1].text = definition
print("Updated and sorted Table 43 (Glossary).")

# 7. Delete Database Management and Functional Instructions elements
body = doc.element.body
elements = list(body)

db_idx = -1
func_idx = -1
faq_idx = -1

for idx, elem in enumerate(elements):
    text = ''
    if elem.tag.endswith('p'):
        p = docx.text.paragraph.Paragraph(elem, doc)
        text = p.text.strip()
    if text == "Database Management & Data Table Operations":
        db_idx = idx
    elif text == "Functional Instructions":
        func_idx = idx
    elif text == "Troubleshooting & FAQs":
        faq_idx = idx

print(f"Dynamic indices found - DB: {db_idx}, Func: {func_idx}, FAQ: {faq_idx}")

# Delete Database Management section first
if db_idx != -1 and func_idx != -1:
    print(f"Deleting Database Management section from element {db_idx} to {func_idx - 1}...")
    for elem in elements[db_idx:func_idx]:
        body.remove(elem)

# Re-list and find new indices
elements = list(body)
func_idx = -1
faq_idx = -1
for idx, elem in enumerate(elements):
    text = ''
    if elem.tag.endswith('p'):
        p = docx.text.paragraph.Paragraph(elem, doc)
        text = p.text.strip()
    if text == "Functional Instructions":
        func_idx = idx
    elif text == "Troubleshooting & FAQs":
        faq_idx = idx

print(f"New dynamic indices found - Func: {func_idx}, FAQ: {faq_idx}")

if func_idx != -1:
    start_del = func_idx + 1
    if faq_idx != -1:
        end_del = faq_idx
    else:
        end_del = len(elements)
    print(f"Deleting Functional Instructions body from element {start_del} to {end_del - 1}...")
    for elem in elements[start_del:end_del]:
        body.remove(elem)
print("Deleted existing sections successfully.")


# 8. Rebuild the Functional Instructions section
p_start = None
for p in doc.paragraphs:
    if p.text == "Functional Instructions":
        p_start = p
        break

if p_start is None:
    raise ValueError("Paragraph 'Functional Instructions' not found!")

# Helper functions for insertion
def insert_p(after_p, text, style='normal'):
    new_p = doc.add_paragraph(text, style=style)
    after_p._p.addnext(new_p._p)
    return new_p

def insert_h1(after_p, text):
    return insert_p(after_p, text, style='Heading 1')

def insert_h2(after_p, text):
    return insert_p(after_p, text, style='Heading 2')

def insert_h3(after_p, text):
    return insert_p(after_p, text, style='Heading 3')

def insert_table(after_p, headers, rows):
    num_cols = len(headers)
    num_rows = len(rows) + 1
    table = doc.add_table(rows=num_rows, cols=num_cols)
    
    # Apply custom styling via XML
    set_table_borders(table)
    
    # Header cells
    for col_idx, h in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_margins(cell)
        
    # Data cells
    for r_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[col_idx]
            cell.text = str(val)
            set_cell_margins(cell)
            
    after_p._p.addnext(table._tbl)
    
    # Add a normal paragraph after table to continue writing text
    new_p = doc.add_paragraph()
    new_p.style = 'normal'
    table._tbl.addnext(new_p._p)
    return new_p

def insert_img(after_p, image_path, page_name):
    new_p = doc.add_paragraph()
    new_p.style = 'normal'
    after_p._p.addnext(new_p._p)
    
    if image_path and os.path.exists(image_path):
        try:
            run = new_p.add_run()
            run.add_picture(image_path, width=docx.shared.Inches(5))
        except Exception as e:
            new_p.text = f"[INSERT SCREENSHOT – {page_name}]"
    else:
        new_p.text = f"[INSERT SCREENSHOT – {page_name}]"
        
    return new_p

# Let's start building!
curr = p_start

# We will write a series of insertions
print("Inserting new rewritten sections...")

# --- ROLE 1: EMPLOYEE / REQUESTOR ---
curr = insert_h1(curr, "Role: Employee / Requestor")
curr = insert_p(curr, "The Employee / Requestor role represents general company staff members who initiate purchase requests (PR1), evaluate proposed substitute items from suppliers, and monitor delivery progress. The primary account used is the employee role account (employee@fortune.com).")

# 1.1 Dashboard
curr = insert_h2(curr, "1.1 Dashboard")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Provides a central overview of requisition statuses, alert notifications, and quick shortcuts.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Welcome Header", "Welcome back, employee!", "Displays logged-in user name"],
    ["KPI Card", "Total Requests", "Total count of requests submitted"],
    ["KPI Card", "Pending Approval", "Count of requests currently in approval workflow"],
    ["KPI Card", "Approved", "Count of successfully approved requests"],
    ["KPI Card", "Rejected", "Count of rejected requests"],
    ["Table", "Recent Requests Table", "Columns: PR1 No., Purpose, Submitted, Priority, Status, View link"],
    ["Button", 'New PR1', "Navigates to creation form"],
    ["Banner", "You have substitute items awaiting decision", "Displays if pending substitutes exist"],
    ["Sidebar Menu", "Dashboard, My Requests, Delivery Status, Substitute Review", "Core module links"],
    ["Header Icons", "Message Center (badge count), Bug Track, Profile Dropdown, Sign out", "Utility actions"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "1. Select Dashboard from the sidebar. Review total pending request counters. If the substitute items warning banner is visible, click on it to navigate to the Substitute Review workspace.")
curr = insert_img(curr, "docs/approver/extracted_media/image7.png", "Employee Dashboard")
curr = insert_p(curr, "Caption: Employee Dashboard - The landing workspace for the employee. It displays four key performance indicator (KPI) metric cards for Total Requests, Pending Approval, Approved, and Rejected, alongside the Recent Requests table which lists the 5 most recent requisitions with columns for PR1 No., Purpose, Date Submitted, Priority, and Status.")
curr = insert_h3(curr, "Status Glossary")
curr = insert_table(curr, ["Status Badge", "Description"], [
    ["Pending Approval", "Yellow status badge indicating that the request is currently routing through approvals."],
    ["Approved", "Green status badge indicating that the request is fully approved."],
    ["Rejected", "Red status badge indicating that the request has been rejected and cannot be edited."]
])
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Successful loading of metrics, tables, and warnings on dashboard landing.")

# 1.2 My Requests
curr = insert_h2(curr, "1.2 My Requests (/pr1)")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Lists all purchase requisitions created by the requisitioner, supporting searching and filtering.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Search Bar", "PR1 number or purpose...", "Input field for text queries"],
    ["Filter Dropdown", "STATUS", "Filters table rows by requisition status"],
    ["Date Pickers", "DATE CREATED FROM / DATE CREATED TO", "Filter rows by creation date range"],
    ["Button", 'Apply Filters', "Submits search and filter parameters"],
    ["Button", 'Clear', "Resets search and filter inputs"],
    ["Button", 'New PR1', "Navigates to creation form"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "1. Type a query in the search input or select status filters, then click the Apply Filters button. The table updates dynamically. Click Clear to reset.")
curr = insert_img(curr, "docs/approver/extracted_media/image49.png", "Requisitions List")
curr = insert_p(curr, "Caption: Purchase Requisitions (PR1) List View - Shows the historical list of all requests submitted by the employee. It includes search inputs for text-based queries on PR1 numbers and purposes, a dropdown to filter by document status, and date pickers to filter by creation ranges.")
curr = insert_h3(curr, "Status Glossary")
curr = insert_table(curr, ["Status Badge", "Description"], [
    ["Draft", "Grey badge showing the request is saved locally but not yet submitted."],
    ["Pending Warehouse Validation", "Amber badge showing the request is awaiting warehouse stock verification."],
    ["PR2 — Pending Approval", "Orange badge indicating Phase 1/Phase 2 approvals are active."],
    ["Canvassing Complete", "Blue badge indicating supplier bidding has closed and winners are awarded."],
    ["Completed (GRN Closed)", "Green badge showing items are received at the warehouse and the request is closed."]
])
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Requisition table correctly updates rows based on filter queries.")

# 1.3 Creating a New PR1
curr = insert_h2(curr, "1.3 Creating a New PR1 (/pr1/new)")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Enables creating and submitting new purchase requisitions.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Read-only Info", "REQUISITIONER / DEPARTMENT / DATE", "Employee / Operations / Today's Date"],
    ["Input Suffix", "PR1-2026- [e.g. 001]", "Custom suffix text input for PR1 sequence"],
    ["Dropdown Select", "PURPOSE *", "Select request business category"],
    ["Date Picker", "DATE REQUIRED *", "Target delivery date required"],
    ["Table", "Items Requested Table", "Columns: Item Code, Description, Unit, SOH, Req. Qty, Raw Mat. Checkbox"],
    ["Button", "Add Item", "Appends a new line item row to the table"],
    ["Button", 'Submit PR1', "Validates and submits the request"],
    ["Button", "Save Draft", "Saves request as draft without submitting"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "Workflow 1: Empty Form Validation\n1. Leave all inputs empty and click Submit PR1. Browser highlights required fields and displays validation bubbles: 'Purpose is required' and 'At least one item is required'.")
curr = insert_img(curr, "docs/approver/extracted_media/image47.png", "New PR1 Empty Form")
curr = insert_p(curr, "Caption: New PR1 Form (Validation State) - The interface displayed when a user attempts to submit an empty form. The browser triggers validation error highlights under the Purpose dropdown and displays a notice stating that at least one item must be added to the items requested table.")
curr = insert_p(curr, "Workflow 2: Submitting a PR1\n1. Enter custom sequence suffix, select a Purpose, choose required date, click Add Item, enter description, unit, and quantity. Check 'Raw Mat.' if applicable. Click Submit PR1.")
curr = insert_img(curr, "docs/approver/extracted_media/image54.png", "Requisition Form Completed")
curr = insert_p(curr, "Caption: New PR1 Form (Populated State) - Illustrates the completed request form. It shows the requisitioner and department metadata fields, custom sequence input (e.g. 001), selected purpose (Operations), required date, and a populated table of requested items with item code, description, unit, and requested quantity.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Form submission creates a new PR1 record, setting status to 'Pending Warehouse Validation'.")

# 1.4 Viewing PR1 Details
curr = insert_h2(curr, "1.4 Viewing PR1 Details (/pr1/{pr1Id})")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Displays detailed information, line items, and approval signatory logs for a specific PR1.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Header Block", "PR1 No. & Status Badges", "Displays document number, priority, and current status"],
    ["Details Panel", "REQUEST HEADER", "Requisitioner, department, purpose, dates, and references"],
    ["Table", "ITEMS REQUESTED", "Columns: #, Item Code, Description, Type, Unit, SOH, Req. Qty, Warehouse Route"],
    ["Timeline", "SIGNATORIES Tracker", "Chronological sign-off steps and active actor indicators"],
    ["Button", "Print", "Opens standard browser print view of the requisition"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "1. Click on a PR1 No. or View link in any list page. Inspect details. Scroll down to the Signatories timeline to see who has reviewed/approved the request.")
curr = insert_img(curr, "docs/approver/extracted_media/image20.png", "Requisition Detail View")
curr = insert_p(curr, "Caption: Requisition Details Page - Displays the details of a specific PR1. The upper portion shows the requisitioner information and item routing decisions, while the bottom section displays the Signatories timeline tracker highlighting approved steps and pending signatories in the workflow.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Detailed request parameters and signature history loaded successfully.")

# 1.5 Substitute Review
curr = insert_h2(curr, "1.5 Substitute Review (/substitutes)")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Lists all requisitions with proposed alternative supplier products.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["KPI Counter", "Pending Decision / Accepted / Rejected", "Departmental substitute decision counters"],
    ["Search Input", "SEARCH", "Filter items by PR1 number, purpose, or supplier"],
    ["Filter Dropdown", "STATUS", "Filter cards by substitute status (All, Decided, Pending)"],
    ["Cards Grid", "Substitute Cards", "Groups items by PR1. Shows description and number of options offered"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "1. Click Substitute Review in the sidebar. Locate card with pending substitutes. Click the card to open side-by-side details.")
curr = insert_img(curr, "docs/approver/extracted_media/image34.png", "Substitute List")
curr = insert_p(curr, "Caption: Substitute Review Dashboard - A queue listing all purchase requests that contain alternative supplier products. Requisitioners can see cards indicating the number of options offered by suppliers and click on them to navigate to the comparison details page.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Substitute review queue and metrics populated.")

# 1.6 Substitute Item Details Comparison
curr = insert_h2(curr, "1.6 Substitute Item Details Comparison (/substitutes/{pr1Id})")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Enables comparing requested items side-by-side with proposed alternatives and submitting decisions.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Left Panel", "YOU REQUESTED", "Displays original requested item name, description, and quantity"],
    ["Right Panel", "SUPPLIER IS OFFERING", "Displays alternative product description, lead time, and remarks"],
    ["Text Field", "Price Hidden Label", "Masks the cost details from general employee view"],
    ["Textarea", "Notes", "Input field to enter justification for decision"],
    ["Button", 'Accept Substitute', "Approves the proposed supplier item"],
    ["Button", 'Reject Substitute', "Declines the proposed item, returning to source request"],
    ["Button", "Change decision to...", "Resets decision state to allow re-selection"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "1. Review comparison cards. Prices display as 'Price hidden'. Enter remarks in the notes box, then click Accept Substitute or Reject Substitute. The status updates immediately. To reverse selection, click the Change decision button.")
curr = insert_img(curr, "docs/approver/extracted_media/image30.png", "Comparison View")
curr = insert_p(curr, "Caption: Substitute Item Comparison - A side-by-side split screen workspace. The left panel shows the original item requested by the employee. The right panel displays the alternative product offered by the supplier (including unit price, lead time, and remarks) with the price hidden from the general employee view to prevent bias.")
curr = insert_h3(curr, "Status Glossary")
curr = insert_table(curr, ["Status Badge", "Description"], [
    ["Pending", "Yellow badge showing that no decision has been taken on the alternative."],
    ["Accepted", "Green badge showing the substitute item is approved for procurement."],
    ["Rejected", "Red badge showing the substitute item is rejected and the system must source the original."]
])
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Decisions are saved, updating downstream canvassing routes.")

# 1.7 Delivery Status
curr = insert_h2(curr, "1.7 Delivery Status (/delivery)")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Tracks shipping milestones and transit logs for approved purchase orders.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Tabs", "All, Pending, Scheduled, In Transit, Delayed, Delivered", "Filters delivery cards by active status"],
    ["Search Input", "SEARCH", "Filter by PO number, purpose, supplier, or warehouse"],
    ["Cards Grid", "Delivery List Cards", "Displays PO number, status badge, supplier, and estimated arrival"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "1. Navigate to Delivery Status. Search by PO number. Click the card chevron to open detail timelines.")
curr = insert_img(curr, "docs/approver/extracted_media/image16.png", "Delivery List")
curr = insert_p(curr, "Caption: Delivery Status Tracker - Lists active deliveries and shipping milestones for approved purchase orders. Requisitioners can click tabs (All, Pending, Scheduled, In Transit, Delayed, Delivered) to filter the list and search for specific PO numbers or suppliers.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Delivery list populates according to status tabs.")

# 1.8 Delivery Details
curr = insert_h2(curr, "1.8 Delivery Details (/delivery/{deliveryId})")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Displays destination parameters and transit history timelines for deliveries.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Information Block", "DELIVERY INFO", "Supplier, Requisitioner, Deliver To, Address"],
    ["Dates Panel", "KEY DATES", "Actual Delivery Date, Estimated Arrival Date"],
    ["Timeline", "STATUS HISTORY", "Chronological log showing status transitions, actors, and remarks"],
    ["Label", "Price hidden", "Masks pricing columns from general employee view"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "1. Open delivery card. Review transit milestones. This screen is read-only for employees; no buttons or actions are present.")
curr = insert_img(curr, "docs/approver/extracted_media/image48.png", "Delivery Details")
curr = insert_p(curr, "Caption: Delivery Details - Displays the read-only shipment details for requisitioners. It maps out supplier information, destination addresses, actual delivery dates, and a chronological history timeline log tracking every status transit step and coordinator note.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Detailed shipment metrics are displayed in a secure, read-only view.")

_helpers = (insert_h2, insert_h3, insert_p, insert_table, insert_img)
curr = build_employee_utilities(curr, _helpers)


# --- ROLE 2: WAREHOUSE ---
curr = insert_h1(curr, "Role: Warehouse")
curr = insert_p(curr, "The Warehouse role handles PR1 stock-on-hand (SOH) validation and processes incoming supplier deliveries by completing Goods Receipt Notes (GRN). The primary account is warehouse@fortune.com (Warehouse Staff position).")

curr = build_warehouse_sections(curr, _helpers)


# --- ROLE 3: PROCUREMENT ---
curr = insert_h1(curr, "Role: Procurement")
curr = insert_p(curr, "The Procurement role acts as the core administrator for purchasing. Responsibilities include creating RFQs, canvassing supplier quotations, awarding winners, generating and routing PR2 memos, issuing Purchase Orders, approving supplier accreditations, and verifying catalog products. The primary account is the procurement officer account (procurement@fortune.com, Procurement Staff position).")

curr = build_procurement_sections(curr, _helpers)


# --- ROLE 4: APPROVER ---
curr = insert_h1(curr, "Role: Approver")
curr = insert_p(curr, "The Approver role represents departmental and organizational decision-makers who hold approval sign-off authority. This includes Supervisor, Department Head, Director, and Finance Director. Accounts sign off on PR1s, PR2s, and POs.")

curr = build_approver_sections(curr, _helpers)


# --- ROLE 5: SUPPLIER ---
curr = insert_h1(curr, "Role: Supplier")
curr = insert_p(curr, "The Supplier Representative manages their accreditation details, submits catalog products, reviews incoming RFQs, inputs pricing bids, acknowledges POs, and schedules shipments. The primary account used is supplier@fortune.com (Supplier Role).")

# 5.1 Dashboard
curr = insert_h2(curr, "5.1 Dashboard (/dashboard)")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Provides a central interface for the supplier to track accreditation status, open RFQs, active POs, and pending delivery shipments.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["KPI Card", "Open RFQs", "Count of active bids open for quoting"],
    ["KPI Card", "Active POs", "Count of POs awaiting acknowledgment or delivery"],
    ["KPI Card", "Pending Deliveries", "Count of shipments in transit or scheduled"],
    ["Status Card", "Accreditation Status", "Displays the current supplier validation status (e.g. Accredited, Pending, Under Review)"],
    ["Panel List", "Recent Notifications", "List of recent system alerts and updates"],
    ["Sidebar Menu", "Dashboard, Accreditation, Product Catalog, Quotations, Purchase Orders, Deliveries", "Core supplier modules"],
    ["Header Icons", "Message Center, Bug Track, Profile Dropdown", "Access chat, file bugs, and configure account settings"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "1. Navigate to the Dashboard. Review the active KPI indicators. If any alerts are visible under notifications, click on them to navigate to the respective module.")
curr = insert_img(curr, "docs/supplier/screenshots/dashboard.png", "Supplier Dashboard")
curr = insert_p(curr, "Caption: Supplier Dashboard - Landing page for supplier portal. It shows KPI cards for Open RFQs, Active POs, and Pending Deliveries, alongside the organization's current Accreditation Status.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Metrics and status badges load dynamically matching supplier database records.")

# 5.2 Accreditation Details
curr = insert_h2(curr, "5.2 Accreditation Details (/supplier/accreditation)")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Allows suppliers to review company profiles and check the status of required compliance certificates.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Header Panel", "Accreditation Details", "Displays Legal Name, Status, Evaluation rating, and Date Accredited"],
    ["Table", "Certification Files", "List of compliance documents (DTI Certificate, Mayor's Permit, Tax Clearance, SEC Registration)"],
    ["Status Badge", "Verified / Pending", "Individual status indicators for each document slot"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "1. Select Accreditation from the sidebar. Inspect your registration profile parameters. Scroll down to review the status of submitted compliance files.")
curr = insert_img(curr, "docs/supplier/screenshots/accreditation.png", "Supplier Accreditation profile")
curr = insert_p(curr, "Caption: Supplier Accreditation Profile - Displays contact details, date accredited, legal name, and evaluation ratings.")
curr = insert_img(curr, "docs/supplier/screenshots/accreditation_docs.png", "Accreditation Documents")
curr = insert_p(curr, "Caption: Accreditation Documents Table - Lists required compliance documents (DTI Certificate, Mayor's Permit, Tax Clearance, SEC Registration) and their verification states.")
curr = insert_h3(curr, "Status Glossary")
curr = insert_table(curr, ["Status Badge", "Description"], [
    ["Accredited", "Green status indicating the supplier is fully approved and can bid on RFQs."],
    ["Pending Review", "Yellow status indicating submitted compliance documents are undergoing check."],
    ["Verified", "Green file badge showing a document was audited and approved by procurement."]
])
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Supplier contact fields and file table load correctly.")

# 5.3 Product Catalog
curr = insert_h2(curr, "5.3 Product Catalog (/supplier/products)")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Enables managing product offerings, adding new items, and checking TSQA verification status.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Button", 'Add Product', "Launches the catalog product creation screen"],
    ["Table", "Product Catalog Table", "Columns: Image, SKU, Category, Product Name, Unit Price, Stock, Status, Action"],
    ["Form Fields", "Product Code, Name, Category, Description, Unit Price, Stock, Image upload slot", "Fields inside the Add Product screen"],
    ["Button", 'Submit Product', "Saves new product details to database"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "Workflow 1: Viewing Product Catalog\n1. Select Product Catalog from the sidebar. Review active SKUs, unit pricing, stock on hand, and TSQA validation status.")
curr = insert_img(curr, "docs/supplier/screenshots/products_list.png", "Product Catalog List")
curr = insert_p(curr, "Caption: Supplier Product Catalog - Displays registered products, categories, SKU codes, unit prices, stock levels, and TSQA verification badges.")
curr = insert_p(curr, "Workflow 2: Adding a Catalog Product\n1. Click the Add Product button. Fill in Code, Name, select a Category, enter Unit Price, and Stock. Click the image upload slot to attach product photos. Click Submit Product.")
curr = insert_img(curr, "docs/supplier/screenshots/product_add.png", "Add Product Form")
curr = insert_p(curr, "Caption: Add Catalog Product - Form to add new products to the catalog, showing product code, name, category, unit price, stock, and photo upload slot.")
curr = insert_p(curr, "Workflow 3: Viewing Product details\n1. Click the View link on a catalog product row to open the detailed product view, including its approval logs.")
curr = insert_img(curr, "docs/supplier/screenshots/product_details.png", "Product Details View")
curr = insert_p(curr, "Caption: Supplier Product Details - Displays product specifications and historical verification decisions (Pending TSQA / Verified).")
curr = insert_h3(curr, "Status Glossary")
curr = insert_table(curr, ["Status Badge", "Description"], [
    ["Verified", "Green badge indicating the product is approved and can be selected for RFQ bids."],
    ["Pending TSQA", "Yellow badge showing that the item sample is awaiting inspection by quality assurance."]
])
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Successful catalog management, insertion, and details review.")

# 5.4 RFQ Quotations
curr = insert_h2(curr, "5.4 RFQ Quotations (/supplier/quotations)")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Lists invitations to bid on corporate requisitions and provides a workspace to submit pricing proposals.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Table", "Quotations Queue Table", "Columns: RFQ Number, Closing Date, Status, Action"],
    ["Link", "View Quote / Edit Quote", "Directs to the RFQ bidding workspace"],
    ["Select Field", "Select Product", "Drop-down selector to link bid line to catalog product"],
    ["Number Input", "Unit Price", "Pricing per unit of measure"],
    ["Number Input", "Lead Time (Days)", "Guaranteed shipping lead time in days"],
    ["Text Input", "Remarks", "Optional remarks for proposed alternatives"],
    ["Button", 'Submit Quotation', "Finalizes the bid and routes to corporate procurement comparison matrix"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "Workflow 1: Viewing RFQ Queue\n1. Select Quotations from the sidebar. Inspect active RFQ numbers and closing deadlines.")
curr = insert_img(curr, "docs/supplier/screenshots/rfq_list.png", "RFQ List Queue")
curr = insert_p(curr, "Caption: Quotations Queue - Lists active Request for Quotation (RFQ) invitations showing closing dates, reference numbers, and active bidding statuses.")
curr = insert_img(curr, "docs/supplier/screenshots/rfq_scrolled.png", "RFQ List Scrolled")
curr = insert_p(curr, "Caption: Quotations Queue (Scrolled) - Displays further quotation details and historical closed bids.")
curr = insert_p(curr, "Workflow 2: Submitting a Bid\n1. Click View Quote on an open RFQ. Review requested items and specifications. Link each requested item to one of your catalog products, input Unit Price and Lead Time, enter optional remarks, and click Submit Quotation.")
curr = insert_img(curr, "docs/supplier/screenshots/rfq_detail.png", "RFQ Bidding workspace")
curr = insert_p(curr, "Caption: RFQ Specifications Header - Shows the RFQ ID, closing date, and company billing/shipping details for the quotation form.")
curr = insert_img(curr, "docs/supplier/screenshots/rfq_lines_details.png", "RFQ Bidding lines")
curr = insert_p(curr, "Caption: RFQ Bid Entry Form - Shows requested items. Suppliers link each line to a product from their catalog and input unit prices, lead times, and remarks.")
curr = insert_img(curr, "docs/supplier/screenshots/rfq_submit.png", "RFQ Submission")
curr = insert_p(curr, "Caption: Quotation Submission Actions - The submission panel. Clicking 'Submit Quotation' registers the supplier's pricing bids in the database.")
curr = insert_h3(curr, "Status Glossary")
curr = insert_table(curr, ["Status Badge", "Description"], [
    ["Active", "Green status showing that the RFQ is currently open and accepting pricing quotes."],
    ["Submitted", "Blue status showing that the supplier has successfully submitted a quotation."],
    ["Closed", "Grey status showing that bidding has ended and no further bids can be accepted."]
])
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Bid submission updates the RFQ state, making quotes visible to procurement agents.")

# 5.5 Purchase Order Acknowledgment
curr = insert_h2(curr, "5.5 Purchase Order Acknowledgment (/supplier/po)")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Provides visibility over issued POs and enables committing to delivery dates.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Table", "Purchase Orders Table", "Columns: PO Number, Date, Total Amount, Status, Action"],
    ["Link", 'View PO', "Opens detail and acknowledgment panel"],
    ["Date Picker", "Commitment Date *", "Required date selector to commit to delivery"],
    ["Textarea", "Remarks", "Optional acknowledgment details"],
    ["Button", 'Acknowledge PO', "Submits signature commitment and issues PO"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "Workflow 1: Reviewing POs\n1. Click Purchase Orders in the sidebar. Locate POs marked 'Pending Acknowledgment'.")
curr = insert_img(curr, "docs/supplier/screenshots/po_list.png", "PO List Queue")
curr = insert_p(curr, "Caption: Supplier Purchase Orders - Lists POs issued to the supplier, showing PO numbers, total amounts, and statuses (Pending Acknowledgment / Acknowledged).")
curr = insert_img(curr, "docs/supplier/screenshots/pos_scrolled.png", "PO List Scrolled")
curr = insert_p(curr, "Caption: Supplier PO List (Scrolled) - Displays historical delivered POs and transaction totals.")
curr = insert_p(curr, "Workflow 2: Acknowledging a PO\n1. Click View PO on a pending order. Verify pricing, payment terms, and items requested. Select a Commitment Date, write optional remarks, and click Acknowledge PO.")
curr = insert_img(curr, "docs/supplier/screenshots/po_detail.png", "PO Detail View")
curr = insert_p(curr, "Caption: PO Detail View (Supplier) - Displays details of an issued PO, showing the items table, billing address, and commercial terms.")
curr = insert_img(curr, "docs/supplier/screenshots/po_detail_scrolled.png", "PO Acknowledgment panel")
curr = insert_p(curr, "Caption: PO Acknowledgment Actions - The commitment panel. Suppliers select a committed delivery date, add remarks, and click 'Acknowledge PO'.")
curr = insert_h3(curr, "Status Glossary")
curr = insert_table(curr, ["Status Badge", "Description"], [
    ["Pending Acknowledgment", "Orange badge indicating the PO was sent to the supplier but has not been acknowledged."],
    ["Acknowledged", "Green badge indicating the supplier accepted the PO and committed to a delivery date."],
    ["Delivered", "Grey badge indicating warehouse receipt of items and closed PO."]
])
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "PO acknowledgment generates an active shipping record in the deliveries module.")

# 5.6 Delivery Management
curr = insert_h2(curr, "5.6 Delivery Management (/supplier/delivery)")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Allows suppliers to schedule dispatches, attach digital receipts, and write transit logs.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Table", "Deliveries Table", "Columns: Delivery Number, PO Reference, Estimated Delivery, Status, Action"],
    ["Link", "View / Update", "Opens the shipping tracker detail page"],
    ["Select Field", "Status Dropdown", "Update status (Scheduled, In Transit, Delayed, Delivered)"],
    ["Date Picker", "Actual Delivery Date", "Date picker for scheduling"],
    ["Upload Slot", "Delivery Receipt / Invoice", "Attachments slots for digital receipts"],
    ["Button", 'Confirm Update', "Submits status update and adds log entry"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "Workflow 1: Reviewing Shipments\n1. Select Deliveries from the sidebar. Inspect statuses of current dispatches.")
curr = insert_img(curr, "docs/supplier/screenshots/deliveries_list.png", "Deliveries List")
curr = insert_p(curr, "Caption: Supplier Deliveries - Lists dispatches showing estimated delivery dates, related PO references, and status badges (Scheduled, In Transit, Delivered).")
curr = insert_p(curr, "Workflow 2: Dispatching a Shipment\n1. Click View on a scheduled shipment. Under Update Status, select In Transit. Select delivery date. Upload digital PDF copies of the Delivery Receipt (DR) and Invoice. Click Confirm Update.")
curr = insert_img(curr, "docs/supplier/screenshots/delivery_detail.png", "Delivery detail tracking")
curr = insert_p(curr, "Caption: Delivery Detail (Supplier) - Displays transit details. Suppliers update status to 'In Transit' or 'Delivered' and upload Delivery Receipt (DR) / Invoice documents.")
curr = insert_img(curr, "docs/supplier/screenshots/delivery_detail_scrolled.png", "Delivery Status updates")
curr = insert_p(curr, "Caption: Delivery Dispatch Status Actions - Dropdown selectors to change shipment states and a chronological transit history log.")
curr = insert_h3(curr, "Status Glossary")
curr = insert_table(curr, ["Status Badge", "Description"], [
    ["Scheduled", "Yellow status indicating the shipment date is set but items have not left the facility."],
    ["In Transit", "Blue status indicating the shipment is currently en route to the warehouse."],
    ["Delivered", "Green status indicating items were successfully delivered to the warehouse."]
])
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Updating status logs updates the delivery timeline immediately for warehouse operators.")

# 5.7 Message Center
curr = insert_h2(curr, "5.7 Message Center")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Enables direct communication with procurement staff and buyers.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Sidebar", "Conversation List", "Shows active chats with preview text"],
    ["Search Input", "User Search", "Find procurement staff to message"],
    ["Thread Panel", "Message Thread", "Type and send messages; view read status"],
    ["Upload", "Attachment", "Attach files (max 10 MB, up to 5 per message)"],
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "1. Click the messages icon in the top header. 2. Use User Search to find a procurement contact. 3. Select or start a conversation. 4. Type your message and press Send.")
curr = insert_img(curr, "docs/supplier/screenshots/messages.png", "Supplier Messages")
curr = insert_p(curr, "Caption: Supplier Messages - Communication interface showing a direct conversation with a procurement staff member.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Message is delivered to the recipient and appears in both users' conversation lists.")

# 5.8 User Profile
curr = insert_h2(curr, "5.8 User Profile")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Enables the supplier representative to manage display credentials and update security passwords.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Input", "Full Name", "Editable display name"],
    ["Read-only", "Role / Company", "Supplier account metadata"],
    ["Input", "Current Password", "Required to change password"],
    ["Input", "New Password / Confirm", "Set a new password"],
    ["Button", "Save Name / Change Password", "Submit profile updates"],
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "1. Click your avatar in the top header. 2. Update your display name and click Save Name, or enter current and new passwords and click Change Password.")
curr = insert_img(curr, "docs/supplier/screenshots/profile.png", "Supplier Profile Settings")
curr = insert_p(curr, "Caption: Supplier Profile - Displays display name credentials and password security settings.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Profile name or password updates are saved; you remain logged in with updated credentials.")

# 5.9 Bug Reporting Tool
curr = insert_h2(curr, "5.9 Bug Reporting Tool")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Enables reporting portal interface bugs to application administrators.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Input", "Title / URL / Severity", "Required bug summary fields"],
    ["Textarea", "Description / Expected Behavior", "Detailed issue explanation"],
    ["Button", "Submit Bug Report", "Sends report to admin notification email"],
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "Workflow 1: Validation\n1. Click Submit without filling fields. Browser highlights required inputs.\n\nWorkflow 2: Submit a Bug\n1. Fill title, URL, severity, description, and expected behavior. Click Submit Bug Report.")
curr = insert_img(curr, "docs/supplier/screenshots/bugtrack.png", "Supplier Bug Form")
curr = insert_p(curr, "Caption: Supplier Bug Report - Form to report technical issues, showing severity, summary, and location inputs.")
curr = insert_img(curr, "docs/supplier/screenshots/bugtrack_scrolled.png", "Supplier Bug Submit")
curr = insert_p(curr, "Caption: Supplier Bug Report (Scrolled) - Displays the submission buttons and error log inputs.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Bug report is saved and an email alert is sent to the configured admin address.")



# --- ROLE 6: TSQA ---
curr = insert_h1(curr, "Role: TSQA")
curr = insert_p(curr, "The TSQA (Technical Services & Quality Assurance) Staff performs Receiving & Storage Evaluations (RSE) on supplier sample products, records test findings, and determines if catalog submissions meet standards.")

# 6.1 RSE Management
curr = insert_h2(curr, "6.1 RSE Management")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Enables TSQA staff to view active evaluations and self-assign tasks.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Tabs", "Created, Assigned, Under Review, Passed, Failed", "Filters evaluation queue by active RSE status"],
    ["Table", "RSE Queue Table", "Columns: RSE Number (RSE-YYYYMM-XXXX), Supplier, Product, Accreditation Doc, Action"],
    ["Button", 'Self-Assign', "Assigns active RSE to current TSQA logged-in user"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "1. Navigate to RSE Queue, click Created tab to see unassigned tasks. Click Self-Assign. Status changes to 'Assigned', transferring the record to your personal worklist.")
curr = insert_img(curr, "docs/tsqa/screenshots/rse_queue.png", "TSQA RSE Queue")
curr = insert_p(curr, "Caption: TSQA RSE Queue - Landing page for TSQA reviewers. It lists active RSE records, categorized by tabs (Created, Assigned, Under Review, Passed, Failed). Reviewers click 'Self-Assign' to assume responsibility.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "RSE moves to Assigned status and appears in your personal worklist.")

# 6.2 Sample Inspection
curr = insert_h2(curr, "6.2 Sample Inspection")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Allows recording laboratory findings and uploading verification reports.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Button", 'Start Review', "Changes status to 'Under Review'"],
    ["Textarea", "Test Findings", "Enter lab measurements, test details, and results"],
    ["Textarea", "Inspection Remarks", "Enter general inspection remarks"],
    ["Upload Slot", "Inspection Report File", "Upload formal test report (PDF/JPG)"],
    ["Button", 'Save Progress', "Saves draft of current test values"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "1. Open assigned RSE. Click Start Review. Enter metrics in Test Findings and remarks in Inspection Remarks. Upload test report PDF. Click Save Progress.")
curr = insert_img(curr, "docs/tsqa/screenshots/rse_inspection.png", "TSQA RSE Inspection")
curr = insert_p(curr, "Caption: TSQA Sample Inspection - The laboratory data entry page. Reviewers click 'Start Review', record test findings and remarks, and upload the formal test report file.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "RSE status changes to Under Review; findings and report are saved.")

# 6.3 Verdict & Product Approval
curr = insert_h2(curr, "6.3 Verdict & Product Approval")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Enables submitting final evaluation decisions on product accreditation.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Button", 'Submit Passed Verdict', "Approves product; updates status to Passed and product to Verified"],
    ["Button", 'Submit Failed Verdict', "Rejects product; requires failure reason remarks"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "1. Open RSE under review. To approve, click Submit Passed Verdict. To reject, click Submit Failed Verdict, enter reasons in the remarks popup, and click Confirm.")
curr = insert_img(curr, "docs/tsqa/screenshots/rse_verdict.png", "TSQA RSE Verdict")
curr = insert_p(curr, "Caption: TSQA RSE Verdict - The verdict panel. Reviewers select 'Submit Passed Verdict' (verifying the product) or 'Submit Failed Verdict' (requiring a failure reason).")
curr = insert_h3(curr, "Status Glossary")
curr = insert_table(curr, ["Status Badge", "Description"], [
    ["Passed", "Product verified; supplier product status set to Verified"],
    ["Failed", "Product rejected; supplier notified"],
])
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Product status updates to Verified or Rejected; supplier and procurement are notified.")


# --- ROLE 7: SUPER ADMIN ---
curr = insert_h1(curr, "Role: Super Admin")
curr = insert_p(curr, "The Super Admin has complete control over system configurations, user directory setups, master data structures, dynamic approval workflows, and module visibilities. The administrative accounts are managed through position-based permissions and all critical administrator operations are tracked in detail by the audit logs. The primary account used is admin@fortune.com.")

# 7.1 Dashboard
curr = insert_h2(curr, "7.1 Dashboard (/admin or /dashboard)")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Monitors system user totals and records recent administrative activity.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["KPI Card", "Total Users", "Total count of active and inactive user accounts"],
    ["KPI Card", "Roles", "Count of roles defined in the system"],
    ["KPI Card", "Positions", "Count of organizational positions"],
    ["KPI Card", "Departments", "Count of company departments"],
    ["KPI Card", "Audit Logs", "Total number of captured activity logs"],
    ["Table", "Recent Activity Table", "List of the 5 most recent audit log records"],
    ["Link", "View All", "Redirects to the full Audit Logs search dashboard"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "1. Access the administrator workspace. The system displays dashboard metric cards showing total system volume. Scroll down to check the list of recent actions. Click on the View All link to load the complete historical audit log grid.")
curr = insert_img(curr, "docs/admin/screenshots/dashboard.png", "Admin Dashboard")
curr = insert_p(curr, "Caption: Admin Dashboard - Operational dashboard displaying KPI counters for total users, roles, positions, departments, and audit logs. The bottom half contains a real-time list of recent activity logs showing timestamps, actions, and references.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Dashboard metrics and recent activity load; View All navigates to the full audit log page.")

# 7.2 User Management
curr = insert_h2(curr, "7.2 User Management (/admin/users)")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Enables provisioning new user accounts, modifying user roles, positions, and departments, resetting account passwords, and deactivating or reactivating access.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Button", "Create User", "Opens user account creation form modal"],
    ["Search Input", "Search", "Filters user records by name, email, or employee ID"],
    ["Dropdown Select", "Role", "Filters user records by role"],
    ["Dropdown Select", "Department", "Filters user records by department"],
    ["Table", "User Directory Table", "Columns: Name/Email, Role, Department, Position, Status badge, Actions"],
    ["Action Link", "Edit", "Opens user assignment edit modal"],
    ["Action Link", "Deactivate", "Opens account deactivation dialog (displays warning status)"],
    ["Action Link", "Reactivate", "Opens account reactivation dialog"],
    ["Action Link", "Reset Password", "Opens password change modal"],
    ["Form Fields", "Full Name, Email, Role, Department, Position, Password", "Create user form inputs"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "Workflow 1: Creating a User Account\n1. Click Create User. Input the user's Full Name, Email, select their Role, Department, and Position from the dropdown lists. Type a strong Password and click Save.\n\nWorkflow 2: Modifying User Assignments\n1. Search for the target user. Click the Edit link next to their name. In the Edit User Assignment modal, select the new Role, Department, or Position, and click Save.\n\nWorkflow 3: Resetting a User Password\n1. Click the Reset Password action link for the user. Enter a new password in the input field, and click Confirm to overwrite their credentials.\n\nWorkflow 4: Deactivating/Reactivating an Account\n1. To disable a user, click the Deactivate link. Review the deactivation warning dialog and click Confirm. To restore access, click Reactivate and confirm in the reactivate dialog.")
curr = insert_img(curr, "docs/admin/screenshots/users.png", "Admin User Management")
curr = insert_p(curr, "Caption: Admin User Management - User management list view. Features include text searching, role and department filtering, a directory list displaying active/inactive statuses, and actions to edit, deactivate, reactivate, or reset passwords.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "User records are created, updated, deactivated, reactivated, or password-reset successfully; changes take effect on next login.")

# 7.3 Roles Management
curr = insert_h2(curr, "7.3 Roles Management (/admin/roles)")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Provides a read-only overview of system roles and shows user counts assigned to each role.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Table", "Roles Table", "Columns: Role Name, Users count badge, Created Date"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "1. Click Roles in the navigation sidebar. Inspect the roles list. Check the Users count badge to see how many accounts are assigned to each role (e.g., admin, employee, warehouse, procurement, approver, supplier, tsqa).")
curr = insert_img(curr, "docs/admin/screenshots/roles.png", "Admin Roles List")
curr = insert_p(curr, "Caption: Roles Management - Displays system roles, showing the specific role names, the count of active users currently associated with each role, and the date each role was created.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Roles list displays with accurate user counts per role (read-only view).")

# 7.4 Positions Management
curr = insert_h2(curr, "7.4 Positions Management (/admin/positions)")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Manages job positions, maps them to system roles, and controls their active state while warning of workflow dependencies.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Button", "Add Position", "Launches the position creation dialog"],
    ["Table", "Positions Table", "Columns: Position Title, Role, Status badge, Created Date, Actions"],
    ["Action Link", "Edit", "Launches the edit position dialog"],
    ["Action Link", "Deactivate", "Opens safety check deactivation dialog"],
    ["Action Link", "Reactivate", "Opens reactivation dialog"],
    ["Form Fields", "Title, Role ID", "Add/Edit position form inputs"],
    ["Warning Banner", "Workflow Usage Notice", "Displays alerts if the position is currently referenced in active approval workflows"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "Workflow 1: Creating a Position\n1. Click Add Position. Enter the Title and select the corresponding Role ID. Review warnings if applicable, and click Save.\n\nWorkflow 2: Editing a Position\n1. Locate the position, click Edit. Modify the Title or change the assigned Role. Click Save. If the position is used in active approval workflows, a warning banner will appear.\n\nWorkflow 3: Deactivating a Position\n1. Click Deactivate. The system checks active user counts and workflow step configurations. If the position is actively used in approval steps, a warning alert will block deactivation until the steps are updated. Click Confirm to finalize deactivation.")
curr = insert_img(curr, "docs/admin/screenshots/positions.png", "Admin Positions List")
curr = insert_p(curr, "Caption: Positions Management - Lists job positions and their associated system roles, showing status badges (Active/Inactive) and options to edit, deactivate, or reactivate positions.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Positions are created, edited, or deactivated; workflow dependency warnings block unsafe deactivations.")

# 7.5 Departments Management
curr = insert_h2(curr, "7.5 Departments Management (/admin/departments)")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Configures organizational departments and their short codes, protecting departments that have active users assigned.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Button", "Add Department", "Opens the department creation dialog"],
    ["Table", "Departments Table", "Columns: Department Name, Code badge, Status badge, Users count badge, Created Date, Actions"],
    ["Action Link", "Edit", "Opens edit department dialog"],
    ["Action Link", "Deactivate", "Opens safety check deactivation dialog"],
    ["Action Link", "Reactivate", "Opens reactivation dialog"],
    ["Form Fields", "Name, Code", "Add/Edit department inputs"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "Workflow 1: Creating a Department\n1. Click Add Department. Enter the Name (e.g. operations) and a short Code (e.g. OPS). Click Save.\n\nWorkflow 2: Editing a Department\n1. Click Edit on the department row. Update the Name or Code fields, and click Save.\n\nWorkflow 3: Deactivating a Department\n1. Click Deactivate. Review the safety dialog which displays the count of active users currently in the department. Click Confirm to disable the department.")
curr = insert_img(curr, "docs/admin/screenshots/departments.png", "Admin Departments List")
curr = insert_p(curr, "Caption: Departments Management - Directory of departments showing department codes, status badges, the number of users assigned to each department, and action buttons.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Departments are created, edited, or deactivated; active-user safety dialog prevents unsafe deactivation.")

# 7.7 Workflow Configuration
curr = insert_h2(curr, "7.6 Workflow Configuration (/admin/workflows)")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Defines and configures dynamic, sequence-based approval pipelines for different document types, validating rules and protecting active steps.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Table", "Workflows Table", "List of workflows: Code, Name, Steps count, Active Instances, Status badge"],
    ["Panel Panel", "Workflow Step Editor", "Displays active instances count, Add Step button, stepper visualization, and step list table"],
    ["Button", "Add Step", "Launches the step creation dialog"],
    ["Action Button", "Edit", "Launches the step edit dialog"],
    ["Action Button", "Delete", "Opens delete step dialog (with active instance check)"],
    ["Form Fields", "Step Order, Step Label, Role, Position, Min Value, Max Value", "Form fields inside the Add/Edit Step modal"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "Workflow 1: Navigating and Selecting a Workflow\n1. Select a workflow from the list (e.g., PR1_APPROVAL). The Workflow Step Editor panel loads, showing the dynamic stepper diagram.\n\nWorkflow 2: Adding an Approval Step\n1. Click Add Step. Select the Step Order number (determines approval sequence). Write a Step Label. Choose the required Role and Position. If the step is threshold-based, enter the Min Value and Max Value range. Click Save. The system validates that step orders are contiguous and do not skip sequences.\n\nWorkflow 3: Deleting an Approval Step\n1. Click Delete on the target step. The system checks if there are any active document instances currently pending on this workflow. If active documents exist, the system blocks the deletion. If safe, click Confirm Delete.")
curr = insert_img(curr, "docs/admin/screenshots/workflows.png", "Admin Workflows Configuration")
curr = insert_p(curr, "Caption: Workflows Configuration - Approval workflow editor. Clicking on a workflow loads its dynamic signatory steps. Stepper configurations include step order sequence, role and position requirements, and threshold values.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Workflow steps are added, edited, or deleted; active-instance guard prevents deleting steps with pending documents.")

# 7.7 Module Visibility
curr = insert_h2(curr, "7.7 Module Visibility (/admin/module-visibility)")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Controls which sidebar links are visible to users. Admins can toggle default visibilities for roles or customize menu visibility per position, including borrowing modules from other roles.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Dropdown Select", "Role", "Select the target role for configuration"],
    ["Dropdown Select", "Position Scope", "Select role default or a specific position"],
    ["Table", "Role Modules Table", "Columns: Visible (Switch toggle), Label, href, module_key"],
    ["Table", "Added Modules from Other Roles", "Columns: Label, Source Role, module_key, Remove action (Only visible in position scope)"],
    ["Button", "Add Module", "Opens add borrowed module modal"],
    ["Button", "Save Changes", "Saves visibility rules to the database"],
    ["Form Fields", "Source Role, Select Modules", "Add modules modal inputs"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "Workflow 1: Toggling Default Role Visibility\n1. Select a Role. Keep Position Scope as 'Role default'. Toggle the Switch icons in the table to show/hide specific links. Click Save Changes.\n\nWorkflow 2: Borrowing Modules for a Position\n1. Select a Role. Select a specific Position from the scope dropdown. Under 'Added Modules from Other Roles', click Add Module. Select the Source Role, check the boxes for the modules you wish to add, and click Add. Review the list of added modules and click Save Changes.")
curr = insert_img(curr, "docs/admin/screenshots/module_visibility.png", "Admin Module Visibility")
curr = insert_p(curr, "Caption: Admin Module Visibility - Visibility controls. Admins select a role and position scope, toggle switches to show/hide standard modules, and borrow pages from other roles using checkable selectors.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Saved visibility rules update the sidebar for the selected role/position on next page load.")

# 7.8 Audit Logs
curr = insert_h2(curr, "7.8 Audit Logs (/admin/audit)")
curr = insert_h3(curr, "Module Purpose")
curr = insert_p(curr, "Provides a searchable and filterable database of system operations, showing actor details, IP addresses, and payload before/after state changes.")
curr = insert_h3(curr, "Module UI Map")
curr = insert_table(curr, ["Component Type", "Element Label / Value", "Description / Action"], [
    ["Input search", "Search", "Filters logs by action, actor email/name, or document reference"],
    ["Dropdown Select", "Document Type", "Filters logs by document type"],
    ["Dropdown Select", "Action", "Filters logs by exact action name"],
    ["Date Pickers", "Date Range", "Filters logs by start and end timestamps"],
    ["Table", "Audit Logs Table", "Columns: Action, Actor, Document Type, Document Reference, IP Address, Timestamp"],
    ["Drawer Panel", "Audit Log Detail", "Displays complete log details, including actor User Agent and database payload JSON tree"]
])
curr = insert_h3(curr, "Step-by-Step Workflows")
curr = insert_p(curr, "1. Access the Audit Logs workspace. Type a query in search or apply filters to update the logs table. Click on a log row to open the Audit Log Detail drawer. Expand the formatted JSON tree to review the exact payload fields changed before and after the action.")
curr = insert_img(curr, "docs/admin/screenshots/audit.png", "Admin Audit Logs")
curr = insert_p(curr, "Caption: Admin Audit Logs - System audit list view. Displays action histories, IP addresses, and timestamps with searching and filtering. Row clicks open details to view exact payload state transitions.")
curr = insert_h3(curr, "Expected Outcome")
curr = insert_p(curr, "Filtered audit log results display; row click opens detail drawer with full payload JSON.")

print("All elements generated. Saving document...")
doc.save("docs/test_full_rewrite.docx")
print("Saved document successfully.")
