import shutil
import subprocess
import os
import docx

# 1. Backup the current file if it exists
if os.path.exists("docs/Procurement Documentation (User Manual).docx"):
    print("Backing up current Procurement Documentation (User Manual).docx...")
    shutil.copyfile("docs/Procurement Documentation (User Manual).docx", "docs/Procurement Documentation (User Manual).docx.backup")

# 2. Copy the original template to Procurement Documentation (User Manual).docx
print("Copying original template to Procurement Documentation (User Manual).docx...")
shutil.copyfile("docs/HRIS Documentation (User Manual) as of Aug 12.docx", "docs/Procurement Documentation (User Manual).docx")

# 3. Run the rewriter script in the docs/approver directory context
print("Running rewrite_user_manual_full.py...")
result = subprocess.run(["python", "docs/approver/rewrite_user_manual_full.py"], cwd="c:\\Users\\Rovick\\Desktop\\project", capture_output=True, text=True)
print("STDOUT:", result.stdout)
print("STDERR:", result.stderr)

# 4. Check if the target sections are present in the resulting document
print("=== Verifying sections in generated document ===")
doc = docx.Document("docs/Procurement Documentation (User Manual).docx")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "Troubleshooting" in text or "Appendices" in text or "Glossary" in text:
        print(f"Paragraph {idx}: '{text}'")
