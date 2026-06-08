import shutil
import subprocess
import docx

# Copy the original template to test_full_rewrite.docx
shutil.copyfile("docs/HRIS Documentation (User Manual) as of Aug 12.docx", "docs/test_full_rewrite.docx")

# Read rewrite_user_manual_full.py
with open("docs/approver/rewrite_user_manual_full.py", "r", encoding="utf-8") as f:
    code = f.read()

# Replace document paths in the code to point to the test file
code_modified = code.replace(
    'doc = docx.Document("docs/Procurement Documentation (User Manual).docx")',
    'doc = docx.Document("docs/test_full_rewrite.docx")'
).replace(
    'doc.save("docs/Procurement Documentation (User Manual).docx")',
    'doc.save("docs/test_full_rewrite.docx")'
)

# Write modified code to temporary runner
temp_runner = "scratch/temp_rewrite_runner.py"
with open(temp_runner, "w", encoding="utf-8") as f:
    f.write(code_modified)

# Run the modified script
print("Running full rewrite logic on fresh template...")
result = subprocess.run(["python", temp_runner], cwd=".", capture_output=True, text=True)
print("STDOUT:", result.stdout)
print("STDERR:", result.stderr)

# Verify if sections exist in the resulting docx
print("=== Verifying test_full_rewrite.docx ===")
doc = docx.Document("docs/test_full_rewrite.docx")
print(f"Total paragraphs: {len(doc.paragraphs)}")
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "Troubleshooting" in text or "Appendices" in text or "Glossary" in text:
        print(f"Found paragraph {idx}: '{text}'")
