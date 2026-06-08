import shutil
import subprocess
import os
import docx

# 1. Copy the original template to test_full_rewrite.docx
print("Copying template...")
shutil.copyfile("docs/HRIS Documentation (User Manual) as of Aug 12.docx", "docs/test_full_rewrite.docx")

# 2. Read rewrite_user_manual_full.py
print("Reading rewriter script...")
with open("docs/approver/rewrite_user_manual_full.py", "r", encoding="utf-8") as f:
    code = f.read()

# 3. Replace paths in code to use test_full_rewrite.docx
code_modified = code.replace(
    'doc = docx.Document("docs/Procurement Documentation (User Manual).docx")',
    'doc = docx.Document("docs/test_full_rewrite.docx")'
).replace(
    'doc.save("docs/Procurement Documentation (User Manual).docx")',
    'doc.save("docs/test_full_rewrite.docx")'
)

# 4. Write modified code to docs/approver/temp_rewrite_runner.py
temp_runner = "docs/approver/temp_rewrite_runner.py"
print("Writing temp runner...")
with open(temp_runner, "w", encoding="utf-8") as f:
    f.write(code_modified)

# 5. Run the temporary runner script
print("Running temp runner...")
result = subprocess.run(["python", "docs/approver/temp_rewrite_runner.py"], cwd="c:\\Users\\Rovick\\Desktop\\project", capture_output=True, text=True)
print("STDOUT:", result.stdout)
print("STDERR:", result.stderr)

# 6. Verify headings in test_full_rewrite.docx
print("=== Verifying sections in generated document ===")
doc = docx.Document("docs/test_full_rewrite.docx")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "Troubleshooting" in text or "Appendices" in text or "Glossary" in text:
        print(f"Paragraph {idx}: '{text}'")

# Clean up temp runner
if os.path.exists(temp_runner):
    os.remove(temp_runner)
