import shutil
import docx
import os

# Copy original template to a test file
shutil.copyfile("docs/HRIS Documentation (User Manual) as of Aug 12.docx", "docs/test_procurement.docx")

doc = docx.Document("docs/test_procurement.docx")
body = doc.element.body
elements = list(body)

db_idx = -1
func_idx = -1
faq_idx = -1

for idx, elem in enumerate(elements):
    text = ''
    if elem.tag.endswith('p'):
        p = docx.text.paragraph.Paragraph(elem, doc)
        text = p.text.strip()
    if text == "Database Management & Data Table Operations":
        db_idx = idx
    elif text == "Functional Instructions":
        func_idx = idx
    elif text == "Troubleshooting & FAQs":
        faq_idx = idx

print(f"Original indices found - DB: {db_idx}, Func: {func_idx}, FAQ: {faq_idx}")

# Delete Database Management section first
if db_idx != -1 and func_idx != -1:
    print(f"Deleting Database Management section from element {db_idx} to {func_idx - 1}...")
    for elem in elements[db_idx:func_idx]:
        body.remove(elem)

# Re-list and find new indices
elements = list(body)
func_idx = -1
faq_idx = -1
for idx, elem in enumerate(elements):
    text = ''
    if elem.tag.endswith('p'):
        p = docx.text.paragraph.Paragraph(elem, doc)
        text = p.text.strip()
    if text == "Functional Instructions":
        func_idx = idx
    elif text == "Troubleshooting & FAQs":
        faq_idx = idx

print(f"New dynamic indices found - Func: {func_idx}, FAQ: {faq_idx}")

if func_idx != -1:
    start_del = func_idx + 1
    if faq_idx != -1:
        end_del = faq_idx
    else:
        end_del = len(elements)
    print(f"Deleting Functional Instructions body from element {start_del} to {end_del - 1}...")
    for elem in elements[start_del:end_del]:
        body.remove(elem)
print("Deleted existing sections successfully.")

# Re-inspect to see what headings are left before saving
print("=== Inspecting before save ===")
elements_before_save = list(body)
for idx, elem in enumerate(elements_before_save):
    if elem.tag.endswith('p'):
        p = docx.text.paragraph.Paragraph(elem, doc)
        text = p.text.strip()
        if "Troubleshooting" in text or "Appendices" in text or "Glossary" in text:
            print(f"Found paragraph {idx}: '{text}'")

doc.save("docs/test_procurement.docx")
print("Saved document.")

# Reload and check
doc2 = docx.Document("docs/test_procurement.docx")
print("=== Inspecting after save and reload ===")
for idx, p in enumerate(doc2.paragraphs):
    text = p.text.strip()
    if "Troubleshooting" in text or "Appendices" in text or "Glossary" in text:
        print(f"Found paragraph {idx}: '{text}'")
