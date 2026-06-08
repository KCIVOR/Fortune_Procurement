import docx

print("=== Simulating Second Pass ===")
doc = docx.Document("docs/test_procurement.docx")
body = doc.element.body
elements = list(body)

db_idx = -1
func_idx = -1
faq_idx = -1

for idx, elem in enumerate(elements):
    text = ''
    if elem.tag.endswith('p'):
        p = docx.text.paragraph.Paragraph(elem, doc)
        text = p.text.strip()
    if text == "Database Management & Data Table Operations":
        db_idx = idx
    elif text == "Functional Instructions":
        func_idx = idx
    elif text == "Troubleshooting & FAQs":
        faq_idx = idx

print(f"Second pass indices - DB: {db_idx}, Func: {func_idx}, FAQ: {faq_idx}")

# Let's inspect the paragraphs in the document
print(f"Total paragraphs in doc: {len(doc.paragraphs)}")
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "Troubleshooting" in text or "Appendices" in text or "Glossary" in text:
         print(f"Found paragraph in doc.paragraphs at index {idx}: '{text}'")
